"""Tests for interview.profile (ephemeral parsing) and the in-memory profile
overlay onto a session's InterviewManager. The candidate's JD/resume must feed
the prompts (as {jd}/{resume}) without ever being written to disk."""

import io

import pytest

from interview.interview_manager import InterviewManager
from interview.profile import CandidateProfile, extract_resume_text
from interview.profile_analyzer import CandidateBrief

BASE_YAML = """
interview:
  title: "测试面试"
interviewer:
  name: "测试面试官"
candidate:
  target_role: "原始岗位"
  background: "原始候选人背景"
positions:
  - name: "后端工程师岗位"
    match_keywords: ["后端"]
    core_competencies: "重点考察技术深度：架构与性能"
    business_questions:
      - "请介绍你自己。"
"""


@pytest.fixture
def config_path(tmp_path):
    path = tmp_path / "interview.yaml"
    path.write_text(BASE_YAML, encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# extract_resume_text
# ---------------------------------------------------------------------------


def test_extract_txt_and_md():
    assert extract_resume_text("resume.txt", "三年后端经验".encode()) == "三年后端经验"
    assert extract_resume_text("resume.md", b"# Resume\nPython") == "# Resume\nPython"


def test_extract_unsupported_ext_raises():
    with pytest.raises(ValueError, match="暂不支持"):
        extract_resume_text("resume.jpg", b"xx")


def test_extract_empty_pdf_raises():
    from pypdf import PdfWriter

    buf = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.write(buf)
    with pytest.raises(ValueError, match="未能从文件中提取到文本"):
        extract_resume_text("resume.pdf", buf.getvalue())


def test_extract_docx():
    from docx import Document

    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("五年 Python 开发经验")
    document.add_paragraph("负责订单系统")
    document.save(buf)
    text = extract_resume_text("resume.docx", buf.getvalue())
    assert "五年 Python 开发经验" in text
    assert "负责订单系统" in text


# ---------------------------------------------------------------------------
# CandidateProfile (ephemeral, in-memory)
# ---------------------------------------------------------------------------


def test_profile_summary_reports_counts_not_content():
    prof = CandidateProfile(target_role="测试工程师", jd_text="负责后端服务", resume_text="三年经验")
    assert prof.summary == {
        "target_role": "测试工程师",
        "jd_chars": 6,
        "resume_chars": 4,
        "has_jd": True,
    }


def test_profile_merge_overwrites_only_nonempty():
    prof = CandidateProfile(target_role="原始", jd_text="旧 JD", resume_text="旧简历")
    merged = prof.merge(jd_text="新 JD 内容")
    assert merged.jd_text == "新 JD 内容"
    assert merged.target_role == "原始"  # untouched field preserved
    assert merged.resume_text == "旧简历"
    # whitespace-only input must not wipe an existing value
    assert prof.merge(target_role="   ").target_role == "原始"


def test_profile_is_empty():
    assert CandidateProfile().is_empty()
    assert not CandidateProfile(jd_text="x").is_empty()


# ---------------------------------------------------------------------------
# InterviewManager.apply_candidate_profile — overlay, no disk writes
# ---------------------------------------------------------------------------


def test_overlay_feeds_system_prompt(config_path):
    manager = InterviewManager(config_path)
    brief = CandidateBrief(
        target_role="测试工程师",
        has_jd=True,
        job_requirements=["后端服务能力"],
        candidate_summary="具备三年相关经验",
        projects=[
            {
                "name": "订单系统",
                "summary": "负责核心链路",
                "source_excerpt": "PROJECT_SOURCE_SENTINEL",
            }
        ],
    )
    manager.apply_candidate_profile(
        CandidateProfile(
            target_role="测试工程师",
            jd_text="JD_RAW_SENTINEL_负责后端服务",
            resume_text="RESUME_RAW_SENTINEL_三年经验",
        ),
        brief,
    )
    system = manager.build_system_prompt()
    assert "测试工程师" in system          # target_role overridden
    assert "后端服务能力" in system         # only condensed job requirements
    assert "具备三年相关经验" in system      # only condensed candidate summary
    assert "JD_RAW_SENTINEL" not in system
    assert "RESUME_RAW_SENTINEL" not in system
    assert "PROJECT_SOURCE_SENTINEL" not in system
    assert "PROJECT_SOURCE_SENTINEL" in brief.planner_context()
    # the JD matches this position ("后端"), so its name + competency are in the block
    assert "后端工程师岗位" in system
    assert "技术深度" in system
    assert "原始候选人背景" not in system    # static config background dropped for a candidate


def test_overlay_unmatched_position_not_in_prompt(config_path):
    # A candidate whose JD matches no configured position must NOT be evaluated
    # against a residual, unrelated position (the 产品经理 → 后端 report bug).
    manager = InterviewManager(config_path)
    manager.apply_candidate_profile(
        CandidateProfile(
            target_role="产品经理", jd_text="负责产品需求分析与用户增长", resume_text="五年产品经验"
        ),
        CandidateBrief(target_role="产品经理", has_jd=True, candidate_summary="产品经验"),
    )
    system = manager.build_system_prompt()
    assert "产品经理" in system
    assert "产品经验" in system
    assert "负责产品需求分析与用户增长" not in system
    assert "后端工程师岗位" not in system    # residual bank position must not leak
    assert "技术深度" not in system          # ...nor its core competencies
    assert "原始候选人背景" not in system     # static background cleared for a candidate


def test_unmatched_position_yields_empty_bank(config_path):
    # No matching position → the planner receives an EMPTY bank and generates from the
    # JD; it must NOT fall back to another position's questions/competencies.
    manager = InterviewManager(config_path)
    manager.apply_candidate_profile(
        CandidateProfile(
            target_role="产品经理", jd_text="负责产品需求分析与用户增长", resume_text="五年产品经验"
        )
    )
    assert manager.matched_question_specs() == []
    assert manager.matched_competencies_text() == ""


def test_overlay_latest_profile_replaces(config_path):
    manager = InterviewManager(config_path)
    manager.apply_candidate_profile(CandidateProfile(jd_text="第一版"))
    manager.apply_candidate_profile(CandidateProfile(jd_text="第二版"))
    system = manager.build_system_prompt()
    assert "第二版" in system
    assert "第一版" not in system


def test_overlay_empty_profile_keeps_target_role(config_path):
    manager = InterviewManager(config_path)
    manager.apply_candidate_profile(CandidateProfile())
    assert manager.config.candidate.target_role == "原始岗位"


def test_overlay_never_writes_to_disk(config_path):
    before = config_path.read_text(encoding="utf-8")
    manager = InterviewManager(config_path)
    manager.apply_candidate_profile(
        CandidateProfile(target_role="X", jd_text="Y", resume_text="Z")
    )
    manager.build_system_prompt()
    assert config_path.read_text(encoding="utf-8") == before
